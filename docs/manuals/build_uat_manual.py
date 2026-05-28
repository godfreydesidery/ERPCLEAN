"""
Generator for the Orbix Engine UAT (User Acceptance Testing) manual.

Run: python docs/manuals/build_uat_manual.py
Output: docs/manuals/Orbix-Engine-UAT-Manual.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "Orbix-Engine-UAT-Manual.docx"

QA_URL = "http://16.170.11.41/"
LOCAL_URL = "http://localhost:8081/"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x12, 0x3A, 0x6E)  # dark navy
    return h


def add_para(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_checklist(doc, items):
    """Acceptance-test checklist rendered as [ ] items."""
    for item in items:
        doc.add_paragraph(f"[ ]  {item}", style="List Bullet")


def add_kv_table(doc, rows, *, col1="Field", col2="Value"):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = col1
    hdr[1].text = col2
    for r in hdr:
        for p in r.paragraphs:
            for run in p.runs:
                run.bold = True
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
    t.allow_autofit = False
    t.columns[0].width = Cm(5.0)
    t.columns[1].width = Cm(10.5)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"NOTE — {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"WARNING — {text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xA0, 0x2C, 0x2C)


def add_pagebreak(doc):
    doc.add_page_break()


def build():
    doc = Document()

    # --- Page margins ----------------------------------------------------
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # --- Cover -----------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Orbix Engine")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x12, 0x3A, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("User Acceptance Testing Manual")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Phase 1 MVP · 2026-05-28")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph("")
    doc.add_paragraph("")

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro.add_run(
        "This manual walks an end user (back-office operator, "
        "accountant, storekeeper or salesperson) through hands-on testing of "
        "every module currently shipped in Orbix Engine. Each section lists "
        "what to do, what to watch for, and an acceptance checklist. Use the "
        "tick boxes to record outcomes; report anything that does not match "
        "the expected behaviour to the project owner."
    ).font.size = Pt(11)

    add_pagebreak(doc)

    # --- Table of contents (manual; Word can regenerate via fields, but
    #     a plain list is more portable for non-Word readers) ------------
    add_heading(doc, "Contents", level=1)
    contents = [
        "1.  Getting started — access, login, roles",
        "2.  Pre-test setup",
        "3.  Catalog — items, UoM, VAT, price lists",
        "4.  Parties — customers and suppliers",
        "5.  Procurement — LPO, GRN, supplier invoices, payments, vendor returns",
        "6.  Sales — quotation, invoice, receipt, customer returns, credit notes",
        "7.  Stock — on-hand, counts, oversell, adjustments",
        "8.  Day cash — open and close the business day",
        "9.  Debt — AR aging, AP dunning, chase notes, write-off",
        "10. Reports — dashboard, sales reports, statements, export",
        "11. Known limitations (out of scope for this UAT)",
        "12. How to report issues",
        "Annex A. Permissions cheat-sheet",
    ]
    for line in contents:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.5)

    add_pagebreak(doc)

    # =====================================================================
    # 1. Getting started
    # =====================================================================
    add_heading(doc, "1. Getting started", level=1)

    add_heading(doc, "1.1 Where to access the system", level=2)
    add_kv_table(
        doc,
        [
            ("QA server (shared)", QA_URL),
            ("Local QA-parity (your machine, when running)", LOCAL_URL),
            ("Health check (any URL)", "<base-url>/actuator/health"),
            ("Swagger / API docs", "<base-url>/swagger-ui.html"),
        ],
    )
    add_para(
        doc,
        "Open the URL in Chrome, Edge or Firefox. The Angular web ERP "
        "loads from the same address as the API — there is no separate "
        "front-end host.",
    )

    add_heading(doc, "1.2 Logging in for the first time", level=2)
    add_steps(
        doc,
        [
            "Go to the URL above. You will be redirected to /login.",
            "Enter the username: rootadmin",
            "Enter the password supplied by the project owner "
            "(see orbix-engine-infra/qa/CREDENTIALS.local.md if you have "
            "repo access).",
            "Click Sign in.",
            "On first login the system prompts you to pick an active branch. "
            "Select HQ (the bootstrapped default).",
        ],
    )
    add_note(
        doc,
        "The QA box was deployed with fresh data on 2026-05-28. The only "
        "user that exists at start is rootadmin. Create per-role users "
        "yourself in section 2 before testing role-specific workflows.",
    )

    add_heading(doc, "1.3 Roles you will use", level=2)
    add_para(
        doc,
        "Orbix Engine ships with the following roles out of the box. Each "
        "role bundles the permissions a real operator needs. Create one "
        "user per role so you can switch personas and confirm permission "
        "gates work correctly.",
    )
    add_kv_table(
        doc,
        [
            ("Role", "What they do"),
            ("ADMIN", "Full access — operates the system itself; do NOT "
                      "use this for feature testing"),
            ("Accountant", "Reviews AR / AP, posts receipts, manages debt, "
                           "reads reports, runs write-offs (dual approval)"),
            ("Salesperson", "Raises quotations, invoices, captures receipts, "
                            "processes customer returns and credit notes"),
            ("Storekeeper / Stock controller",
             "Receives GRNs, runs stock counts, authorises oversell, "
             "processes vendor returns"),
            ("Procurement officer",
             "Raises and approves LPOs, posts supplier invoices"),
            ("Cashier", "Reserved for the POS desktop app (not in this UAT)"),
        ],
        col1="Role", col2="What they do",
    )

    add_pagebreak(doc)

    # =====================================================================
    # 2. Pre-test setup
    # =====================================================================
    add_heading(doc, "2. Pre-test setup", level=1)

    add_heading(doc, "2.1 Create users for each role", level=2)
    add_steps(
        doc,
        [
            "Sign in as rootadmin.",
            "Open the side menu — Admin > Users.",
            "Click Create user. Fill in: username, full name, e-mail, "
            "initial password.",
            "Assign the role from the role drop-down. Tick company-wide if "
            "the user should not be restricted to one branch.",
            "Save. Repeat for at least one user per role from the table in "
            "section 1.3.",
        ],
    )
    add_warning(
        doc,
        "Set a different password for each test user and write it down "
        "where the team can find it. The initial password is plain text "
        "and the user will be prompted to change it on first login.",
    )

    add_heading(doc, "2.2 Pick the active branch", level=2)
    add_para(
        doc,
        "Every transactional document (invoice, GRN, sales receipt, till "
        "session, debt note) is stamped with the branch that posted it. "
        "Use the branch selector in the top bar to switch context if you "
        "need to test cross-branch behaviour. For most UAT tests the "
        "bootstrapped HQ branch is enough.",
    )

    add_heading(doc, "2.3 Confirm system health", level=2)
    add_checklist(
        doc,
        [
            "The URL loads and shows the Orbix login screen.",
            "/actuator/health returns {\"status\":\"UP\"}.",
            "/swagger-ui.html lists at least the dashboard, sales-invoice, "
            "supplier-invoice, debt, and reports controllers.",
            "Sidebar shows: Dashboard, Catalog, Parties, Procurement, "
            "Sales, Stock, Debt, Reports, Admin.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 3. Catalog
    # =====================================================================
    add_heading(doc, "3. Catalog", level=1)

    add_heading(doc, "3.1 What's covered", level=2)
    add_bullets(
        doc,
        [
            "Units of measure (UoM) — Each, Carton, Kg, etc.",
            "VAT groups — Standard (18%), Zero-rated (0%), Exempt.",
            "Item groups — hierarchical categories.",
            "Items — sellable goods or materials.",
            "Price lists — wholesale, retail, VIP.",
            "Item barcodes — multiple per item.",
        ],
    )

    add_heading(doc, "3.2 Reference data (do this first)", level=2)
    add_steps(
        doc,
        [
            "Open Catalog > Units of measure. Confirm Each (EACH) and Kg "
            "exist. If not, create them.",
            "Open Catalog > VAT groups. Confirm a Standard VAT group exists "
            "(typically 18% in Tanzania). Create one if missing.",
            "Open Catalog > Item groups. Create at least two groups, e.g. "
            "Beverages and Groceries.",
            "Open Catalog > Price lists. Confirm a Default price list exists.",
        ],
    )

    add_heading(doc, "3.3 Create an item", level=2)
    add_steps(
        doc,
        [
            "Catalog > Items > Create item.",
            "Fill in: code (auto-generated), name, item group, base UoM, "
            "VAT group, default cost and selling price.",
            "Optionally add a barcode under the Barcodes panel.",
            "Save. The item should now appear in the items list with "
            "status ACTIVE.",
        ],
    )

    add_heading(doc, "3.4 Search items (typeahead)", level=2)
    add_para(
        doc,
        "Wherever the UI lets you select an item (sales-invoice line, "
        "vendor-return line, stock-count add line) you should see a "
        "TYPEAHEAD — type 2-3 characters of code or name and pick from "
        "the drop-down. You should NEVER need to type a raw uid or id.",
    )

    add_heading(doc, "3.5 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "I can create an item with code, name, UoM, VAT group.",
            "The item appears in the items list immediately.",
            "I can search the item by code AND by name from the typeahead.",
            "I can archive an item (status flips to ARCHIVED).",
            "I can re-activate an archived item.",
            "I can add a second barcode to an item.",
            "Bulk price list assignments (if exposed) work without errors.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 4. Parties
    # =====================================================================
    add_heading(doc, "4. Parties (customers and suppliers)", level=1)

    add_heading(doc, "4.1 Create a customer", level=2)
    add_steps(
        doc,
        [
            "Parties > Customers > Create customer.",
            "Fill in: name, optional TIN, phone, address, credit limit "
            "(e.g. TZS 500,000), payment terms (e.g. 30 days), default "
            "price list.",
            "Save. The system creates the party record AND the customer "
            "role in one transaction and emits PartyCreated.v1 + "
            "CustomerCreated.v1 outbox events.",
        ],
    )

    add_heading(doc, "4.2 Create a supplier", level=2)
    add_steps(
        doc,
        [
            "Parties > Suppliers > Create supplier.",
            "Fill in: name, TIN, contact, default currency, payment terms "
            "(e.g. 30 days), credit limit-from-supplier (how much they "
            "extend to us).",
            "Save.",
        ],
    )

    add_heading(doc, "4.3 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "I can create a customer with credit-limit and price-list.",
            "I can create a supplier with payment-terms-days.",
            "The party list shows both customers and suppliers (by tab or "
            "by filter).",
            "I can edit a customer's credit limit from the debt module's "
            "customer drill-down (Debt > AR tab > pick customer).",
            "Archiving a party sets status to ARCHIVED without breaking "
            "historic documents that reference it.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 5. Procurement
    # =====================================================================
    add_heading(doc, "5. Procurement", level=1)
    add_para(
        doc,
        "Tests cover the full payable cycle: LPO -> GRN -> supplier "
        "invoice -> supplier payment -> (optional) vendor return -> "
        "vendor credit note -> apply credit note. This is the AP-side "
        "mirror of the sales workflow (section 6).",
    )

    add_heading(doc, "5.1 Raise an LPO", level=2)
    add_steps(
        doc,
        [
            "Procurement > LPOs > Create LPO.",
            "Pick a supplier (typeahead), expected delivery date, and at "
            "least one line (item + uom + qty + unit cost).",
            "Save as Draft. Status should be DRAFT.",
            "Click Submit. If amount > the auto-approval threshold, "
            "status moves to PENDING_APPROVAL.",
            "Sign in as an ADMIN or procurement-approver and click Approve. "
            "Status becomes APPROVED.",
        ],
    )

    add_heading(doc, "5.2 Receive against an LPO (GRN)", level=2)
    add_steps(
        doc,
        [
            "Procurement > GRNs > Create GRN from LPO.",
            "Pick the approved LPO. Lines pre-fill from the LPO lines.",
            "Adjust received quantities (full receipt, short receipt, or "
            "extra by exception).",
            "Save and Post. Stock-IN movements are created and item "
            "on-hand goes up.",
        ],
    )

    add_heading(doc, "5.3 Match supplier invoice", level=2)
    add_steps(
        doc,
        [
            "Procurement > Supplier invoices > Create.",
            "Pick supplier and link one or more GRNs.",
            "System auto-fills lines from the GRN(s). Adjust unit cost if "
            "the supplier billed differently within tolerance.",
            "Post the invoice. AR-equivalent for AP — increments the "
            "supplier's outstanding payable.",
        ],
    )

    add_heading(doc, "5.4 Pay the supplier", level=2)
    add_steps(
        doc,
        [
            "Cash > Supplier payments > Create.",
            "Pick supplier and one or more open supplier invoices.",
            "Enter the amount paid. Allocate across the picked invoices.",
            "Post. The supplier invoice's paidAmount goes up.",
        ],
    )

    add_heading(doc, "5.5 Return goods to vendor (Slice H.1)", level=2)
    add_steps(
        doc,
        [
            "Procurement > Vendor returns > New.",
            "Pick the supplier (TYPEAHEAD — no raw ids). The form refreshes.",
            "Optionally pick a POSTED GRN via the GRN picker — lines "
            "auto-fill from the GRN. Or leave GRN blank and add lines "
            "manually (item typeahead + UoM dropdown + VAT-group dropdown).",
            "Choose reason: DAMAGED / WRONG_ITEM / EXPIRED / OTHER. Tick "
            "Restock if the goods are returning to your shelf; untick if "
            "they are being discarded as damage.",
            "Save and Post. A stock-OUT movement is posted (RETURN_OUT for "
            "restock, DAMAGE for non-restock) and VendorReturnPosted.v1 is "
            "emitted to the outbox.",
            "From the posted return, click Issue credit note. The credit "
            "note is created in POSTED status with allocatedAmount = 0.",
        ],
    )

    add_heading(doc, "5.6 Apply a vendor credit note (Slice H.1)", level=2)
    add_steps(
        doc,
        [
            "Procurement > Vendor returns > Credit notes tab.",
            "Find the credit note row. Click Apply.",
            "The modal lists open supplier invoices for the same supplier. "
            "Pick one and enter the amount to apply (defaults to the lesser "
            "of credit available and invoice outstanding).",
            "Submit. The credit note status flips to PARTIALLY_ALLOCATED "
            "or FULLY_ALLOCATED, and the chosen supplier invoice's "
            "paidAmount goes up. VendorCreditNoteApplied.v1 is emitted.",
        ],
    )

    add_heading(doc, "5.7 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "I can complete LPO -> GRN -> supplier invoice in one session.",
            "Above-threshold LPOs require a second user to approve.",
            "Below-threshold LPOs auto-post if the same user holds the "
            "approve permission.",
            "Cancelling an LPO that has a GRN against it is BLOCKED (400).",
            "Vendor return form uses pickers — no raw uid fields anywhere.",
            "Posting a vendor return reduces stock on-hand by the qty "
            "returned (confirm via Stock > Stock card).",
            "Applying a credit note reduces the supplier invoice's "
            "outstanding by exactly the applied amount.",
            "Trying to over-apply a credit note returns 422 with a clear "
            "message.",
            "Trying to apply a credit note to a different supplier's "
            "invoice returns 422.",
            "Trying to apply a credit note that is already FULLY_ALLOCATED "
            "returns 409.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 6. Sales
    # =====================================================================
    add_heading(doc, "6. Sales", level=1)

    add_heading(doc, "6.1 Raise a sales invoice", level=2)
    add_steps(
        doc,
        [
            "Sales > Invoices > Create invoice.",
            "Pick a customer. Add one or more item lines.",
            "Save as Draft, then Post. If posting would exceed the "
            "customer's credit limit, the post is BLOCKED — try this on "
            "purpose with a customer whose limit is 0.",
            "If you hold SALES_INVOICE.OVERRIDE_CREDIT, the over-limit "
            "post is allowed but a reason must be captured.",
        ],
    )

    add_heading(doc, "6.2 Capture a receipt and allocate", level=2)
    add_steps(
        doc,
        [
            "Sales > Receipts > Create receipt.",
            "Pick the customer and the open invoices. Enter the amount "
            "received and allocate across invoices.",
            "Post. The invoice paidAmount goes up; if equal to "
            "totalAmount, status flips to PAID.",
        ],
    )

    add_heading(doc, "6.3 Process a customer return (Slice C / V30)", level=2)
    add_steps(
        doc,
        [
            "Sales > Returns > New.",
            "Pick a posted sales invoice; lines pre-fill.",
            "Adjust returned quantities and reason. Save and Post.",
            "Restock = true posts a RETURN_IN stock move (on-hand goes "
            "back up); restock = false posts a DAMAGE move.",
            "From the posted return, click Issue credit note. A "
            "customer credit note appears in POSTED status.",
        ],
    )

    add_heading(doc, "6.4 Apply a customer credit note (Slice H)", level=2)
    add_steps(
        doc,
        [
            "Sales > Returns > Credit notes tab.",
            "Click Apply on a credit-note row.",
            "Modal lists the same customer's open invoices. Pick one and "
            "enter amount.",
            "Submit. Status flips to PARTIALLY_ALLOCATED or "
            "FULLY_ALLOCATED; target invoice's paidAmount goes up; if "
            "fully paid, invoice flips to PAID. CustomerCreditNoteApplied"
            ".v1 is emitted.",
        ],
    )

    add_heading(doc, "6.5 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "I can create an invoice with multiple lines and discounts.",
            "Posting over credit limit is blocked unless I have the "
            "override permission.",
            "Allocating a receipt across two invoices updates BOTH "
            "paidAmount values correctly.",
            "Voiding a same-day posted invoice reverses the stock move.",
            "Customer return DRAFT can be cancelled; POSTED cannot.",
            "Credit-note apply enforces same-customer (cross-customer "
            "returns 422).",
            "Credit-note apply enforces the row cap (over-apply 422; "
            "fully-allocated 409).",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 7. Stock
    # =====================================================================
    add_heading(doc, "7. Stock", level=1)

    add_heading(doc, "7.1 Stock on hand and stock card", level=2)
    add_steps(
        doc,
        [
            "Stock > On hand. Filter by branch and item.",
            "Click an item to open the stock card — a chronological view "
            "of every IN / OUT / ADJUST movement with running balance.",
        ],
    )

    add_heading(doc, "7.2 Adjustment", level=2)
    add_steps(
        doc,
        [
            "Stock > Adjustments > New. Pick item, qty (signed), reason.",
            "Above the adjustment threshold (default TZS 50,000 value) the "
            "post requires a supervisor PIN or a second-user approve.",
        ],
    )

    add_heading(doc, "7.3 Stock count", level=2)
    add_steps(
        doc,
        [
            "Stock > Counts > Start count. Pick branch, optionally filter "
            "by item group.",
            "Capture counted quantities per row (variance is shown live).",
            "Click Close. Variances post as STOCK_VARIANCE moves.",
            "If your role lacks STOCK.COUNT.APPROVE, closing requires a "
            "second user.",
        ],
    )

    add_heading(doc, "7.4 Oversell guard", level=2)
    add_steps(
        doc,
        [
            "Try to post a sales invoice for an item whose on-hand is less "
            "than the qty sold. By default it is BLOCKED.",
            "A user with the OVERSELL_OVERRIDE permission can authorise "
            "the post — a modal appears asking for a reason.",
            "After authorising, stock on-hand goes negative. The Negative "
            "stock report (Stock > Reports) lists the item.",
        ],
    )

    add_heading(doc, "7.5 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "Stock card shows correct running balance after a mix of GRN, "
            "sale, adjustment, return.",
            "Adjustment over threshold requires dual control.",
            "Stock count closes cleanly with variances posted.",
            "Oversell is blocked for users without override.",
            "Override succeeds for users with the override permission and "
            "captures a reason.",
            "Negative-stock report lists overridden items.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 8. Day cash
    # =====================================================================
    add_heading(doc, "8. Day cash (open and close business day)", level=1)

    add_heading(doc, "8.1 Open the business day", level=2)
    add_steps(
        doc,
        [
            "Day > Open day. Pick branch and date.",
            "Confirm. Status moves to OPEN.",
        ],
    )

    add_heading(doc, "8.2 Close the business day", level=2)
    add_steps(
        doc,
        [
            "Day > Close day.",
            "The system shows totals for the day (sales, purchases, cash "
            "in, cash out).",
            "Confirm to close. Status moves to CLOSED.",
            "If your role lacks DAY.OVERRIDE you cannot post into a closed "
            "day.",
        ],
    )

    add_heading(doc, "8.3 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "Cannot post invoices or GRNs before the day is OPEN.",
            "Close shows correct totals (verify against Reports > Sales "
            "summary).",
            "Closed day blocks new postings unless DAY.OVERRIDE is held.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 9. Debt
    # =====================================================================
    add_heading(doc, "9. Debt", level=1)

    add_heading(doc, "9.1 AR aging and customer drill-down (Slice G)", level=2)
    add_steps(
        doc,
        [
            "Sign in as accountant.",
            "Debt > AR tab. See 5-bucket aging totals (CURRENT, 1-30, "
            "31-60, 61-90, 90+) and dunning queue sorted by oldest overdue.",
            "Click a row to open /debt/customer/uid/:uid. See header (with "
            "credit limit), aging row, open invoices, recent receipts, "
            "and chase-notes activity log.",
        ],
    )

    add_heading(doc, "9.2 Append and archive a chase note (Slice G)", level=2)
    add_steps(
        doc,
        [
            "On the customer drill-down, type a chase note in the activity "
            "log box and click Add note.",
            "The note appears immediately at the top of the log. The DB "
            "emits PartyNoteCreated.v1 with kind = AR_CHASE.",
            "Click Archive on any note. The note disappears (or shows as "
            "archived). PartyNoteArchived.v1 is emitted.",
        ],
    )

    add_heading(doc, "9.3 AP dunning and supplier drill-down (Slice G.1)", level=2)
    add_steps(
        doc,
        [
            "Debt > AP tab. Same 5-bucket aging totals for supplier "
            "payables, dunning queue sorted by oldest overdue.",
            "Click a supplier row to open /debt/supplier/uid/:uid. See "
            "header (with payment-terms-days), aging row, open AP "
            "invoices, recent supplier payments, AP chase notes.",
            "Append an AP_CHASE note the same way as the AR one.",
        ],
    )

    add_heading(doc, "9.4 Write off a debt with dual approval (Slice G.2)", level=2)
    add_steps(
        doc,
        [
            "From the customer drill-down (AR) or supplier drill-down (AP), "
            "click Write off on an open invoice row.",
            "The modal pre-fills with the invoice's outstanding amount. "
            "Enter a reason. Submit.",
            "If amount is at or below the auto-approval threshold "
            "(default TZS 100,000) and you hold DEBT.WRITE_OFF.APPROVE, "
            "the write-off auto-posts. DebtWriteOffPosted.v1 is emitted "
            "and the invoice falls out of the aging.",
            "If above threshold, status is PENDING_APPROVAL. Sign in as a "
            "DIFFERENT user with DEBT.WRITE_OFF.APPROVE and approve from "
            "the Debt > Write-offs queue. Status flips to POSTED.",
            "Try to approve a pending write-off as the SAME user who "
            "submitted it — the system returns 409 with a Different user "
            "required message.",
        ],
    )

    add_heading(doc, "9.5 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "AR aging totals match the sum of open customer invoices.",
            "AP aging totals match the sum of open supplier invoices.",
            "Chase notes survive a page refresh.",
            "Below-threshold write-off auto-posts in one click.",
            "Above-threshold write-off requires a second user to approve.",
            "Same-user-approve on an above-threshold request is rejected "
            "with 409.",
            "Approved write-off removes the invoice from the aging tables.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 10. Reports
    # =====================================================================
    add_heading(doc, "10. Reports", level=1)

    add_heading(doc, "10.1 Dashboard (Slice F)", level=2)
    add_para(
        doc,
        "Open Dashboard. The home page shows KPI tiles (sales today, "
        "purchases today, LPOs pending approval, debt aging, negative "
        "stock). Each tile drills through to the underlying report.",
    )

    add_heading(doc, "10.2 Sales daily (US-RPT-001)", level=2)
    add_steps(
        doc,
        [
            "Reports > Sales daily.",
            "Filters: branch, business date, optional agent.",
            "Columns: invoice no, customer, total, payment terms.",
            "Click the Export [PDF] [Excel] [CSV] button group.",
            "Verify a file downloads (filename includes the report name "
            "and date).",
        ],
    )

    add_heading(doc, "10.3 Sales summary (US-RPT-002, Slice I)", level=2)
    add_steps(
        doc,
        [
            "Reports > Sales summary.",
            "Pick a business date that has activity.",
            "Verify KPI tiles + sub-tables render: top items, top customers, "
            "payment-method breakdown.",
            "Export to all three formats.",
        ],
    )

    add_heading(doc, "10.4 Z-history (US-RPT-003, Slice I)", level=2)
    add_steps(
        doc,
        [
            "Reports > Z-history.",
            "Pick a date range and optionally a till.",
            "Verify per-till-session rows render with declared cash and "
            "variance columns.",
            "Click a row to drill into the underlying Z-report.",
            "Export.",
        ],
    )

    add_heading(doc, "10.5 Statements (US-RPT-007)", level=2)
    add_steps(
        doc,
        [
            "Reports > Customer statement. Pick a customer and date range. "
            "Verify opening balance, period entries, running balance, "
            "closing balance.",
            "Reports > Supplier statement. Same shape on the AP side.",
            "Export to PDF and confirm the layout is readable.",
        ],
    )

    add_heading(doc, "10.6 Acceptance checklist", level=2)
    add_checklist(
        doc,
        [
            "Every report listed above renders the four standard states "
            "correctly: loading, empty, error, populated.",
            "Export to CSV: file opens in Notepad and Excel without "
            "corruption; UTF-8 BOM present.",
            "Export to Excel: file opens in Excel with headers and data "
            "rows; numeric columns are numeric, dates are dates.",
            "Export to PDF: file opens in any PDF reader; landscape if "
            "more than 6 columns; title and subtitle present.",
            "Trying to export when the table is empty disables the buttons "
            "with a tooltip.",
            "Row-cap message appears (toast) when trying to export more "
            "than 5000 rows.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 11. Known limitations
    # =====================================================================
    add_heading(doc, "11. Known limitations (out of scope for this UAT)", level=1)
    add_para(
        doc,
        "The following items are NOT yet shipped. Do not raise issues "
        "against them — they are tracked in the project plan and will "
        "land in future releases.",
    )
    add_bullets(
        doc,
        [
            "POS desktop app (cashier till) — Flutter Windows app, "
            "Epic 7. Not yet available.",
            "WMS field-sales Android app — Epic 8. Not yet available.",
            "Production module (BOMs, batches, conversions) — Epic 10. "
            "Phase 2.",
            "Scheduled report email — Phase 2.",
            "Background-job export for results above 5000 rows — Phase 2.",
            "Phase 1.1 supermarket additions: sections per branch, "
            "foreign-currency tender per till, layby + pre-orders, "
            "gift cards, weighed items / embedded-weight EAN-13, batch / "
            "expiry tracking, section P&L reports.",
            "Multi-company consolidation, inter-branch transfers, "
            "accounting export, mobile-money gateway, fiscal-printer "
            "driver matrix, iOS WMS, AI features — Phase 3.",
            "Cash refund of a customer or vendor credit note — separate "
            "slice.",
            "Auto-apply credit note against oldest overdue invoice — "
            "follow-up.",
            "Reverse / un-apply an allocation — separate slice.",
        ],
    )

    add_pagebreak(doc)

    # =====================================================================
    # 12. How to report issues
    # =====================================================================
    add_heading(doc, "12. How to report issues", level=1)
    add_para(
        doc,
        "For each defect or unexpected behaviour, capture as much of the "
        "following as you can. The minimum is everything in BOLD.",
    )
    add_bullets(
        doc,
        [
            "Date and time (Africa/Dar_es_Salaam timezone)",
            "Which environment (QA URL or local)",
            "Which user (username + role)",
            "Active branch",
            "Page / URL where the issue occurred",
            "Steps to reproduce (numbered)",
            "Expected behaviour (what the user manual says)",
            "Actual behaviour (what you saw)",
            "Screenshot if possible",
            "Browser console output if available "
            "(F12 -> Console)",
            "Network response — open the failing request in F12 -> "
            "Network -> click the row -> Response tab. Capture HTTP "
            "status code and body.",
        ],
    )

    add_para(doc, "")
    add_para(doc, "Severity guidance:", bold=True)
    add_kv_table(
        doc,
        [
            ("S1 - Blocker", "Cannot continue testing this module."),
            ("S2 - Major", "Workflow broken but a workaround exists."),
            ("S3 - Minor", "Cosmetic, copy / colour / spacing."),
            ("S4 - Suggestion", "Nice-to-have improvement."),
        ],
        col1="Severity", col2="Definition",
    )

    add_pagebreak(doc)

    # =====================================================================
    # Annex A. Permissions cheat-sheet
    # =====================================================================
    add_heading(doc, "Annex A. Permissions cheat-sheet", level=1)
    add_para(
        doc,
        "The names below are the permission codes the backend checks. "
        "When testing 403 / Permission required paths, you can map a "
        "user's missing permission to one of these.",
    )
    add_kv_table(
        doc,
        [
            ("Permission", "Granted to (default)"),
            ("DEBT.READ", "Accountant"),
            ("DEBT.NOTE.CREATE", "Accountant"),
            ("DEBT.NOTE.ARCHIVE", "Accountant"),
            ("DEBT.CREDIT_LIMIT.UPDATE", "Accountant"),
            ("DEBT.WRITE_OFF.REQUEST", "Accountant"),
            ("DEBT.WRITE_OFF.APPROVE", "Accountant (a different user "
                                       "than the requester for >threshold)"),
            ("SALES.MANAGE_INVOICE", "Salesperson"),
            ("SALES.MANAGE_RECEIPT", "Salesperson"),
            ("SALES.MANAGE_RETURN", "Salesperson"),
            ("SALES_INVOICE.OVERRIDE_CREDIT", "Sales manager"),
            ("PROCUREMENT.MANAGE_LPO", "Procurement officer"),
            ("PROCUREMENT.APPROVE_LPO", "Procurement manager"),
            ("PROCUREMENT.MANAGE_INVOICE", "Accountant / procurement officer"),
            ("PROCUREMENT.MANAGE_RETURN", "Procurement officer / accountant"),
            ("STOCK.COUNT.APPROVE", "Stock controller"),
            ("OVERSELL.OVERRIDE", "Store manager"),
            ("DAY.OPEN", "Store manager"),
            ("DAY.CLOSE", "Store manager"),
            ("DAY.OVERRIDE", "Accountant"),
        ],
        col1="Permission", col2="Granted to (default)",
    )

    add_para(doc, "")
    add_para(
        doc,
        "End of manual. Thank you for testing Orbix Engine.",
        italic=True,
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"OK -> {path}")
    print(f"size: {path.stat().st_size:,} bytes")
